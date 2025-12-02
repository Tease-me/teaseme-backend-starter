import io
import logging
from typing import List, Dict
from PyPDF2 import PdfReader
from docx import Document

log = logging.getLogger("document_parser")

async def extract_text_from_pdf(file_obj: io.BytesIO) -> str:
    """Extract text from PDF file"""
    try:
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        log.error(f"PDF extraction error: {e}", exc_info=True)
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

async def extract_text_from_docx(file_obj: io.BytesIO) -> str:
    """Extract text from Word document, including tables"""
    try:
        file_obj.seek(0)
        doc = Document(file_obj)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables (CVs often use tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            log.warning("No text extracted from DOCX file - file might be empty or contain only images")
        
        log.info(f"Extracted {len(text)} characters from DOCX ({len(text_parts)} text parts)")
        return text.strip()
    except Exception as e:
        log.error(f"DOCX extraction error: {e}", exc_info=True)
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, any]]:
    """
    Split text into overlapping chunks with metadata.
    Returns list of dicts with 'content' and 'metadata' keys.
    
    Args:
        text: The text to chunk
        chunk_size: Target number of words per chunk
        overlap: Number of words to overlap between chunks
    
    Returns:
        List of dicts with 'content' and 'metadata' keys
    """
    if not text or not text.strip():
        return []
    
    words = text.split()
    chunks = []
    chunk_index = 0
    
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunk_text = " ".join(chunk_words)
        
        if chunk_text.strip():
            chunks.append({
                "content": chunk_text,
                "metadata": {
                    "chunk_index": chunk_index,
                    "start_word": i,
                    "end_word": min(i + chunk_size, len(words)),
                    "total_words": len(words)
                }
            })
            chunk_index += 1
        
        # Move forward, accounting for overlap
        i += chunk_size - overlap
        if i >= len(words):
            break
    
    log.info(f"Chunked text into {len(chunks)} chunks (chunk_size={chunk_size}, overlap={overlap})")
    return chunks

